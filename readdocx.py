import docx
import json
import argparse
import sys

# Настройка парсера аргументов командной строки
parser = argparse.ArgumentParser(description="Extract text and tables from a DOCX file.")
parser.add_argument("docx_path", type=str, help="Path to the DOCX file")
args = parser.parse_args()
docx_path = args.docx_path  # Читаем путь из аргумента командной строки

# Список для хранения данных всех таблиц и текстов
data = {"tables": [], "text": "", "texttables": []}

def extract_text(doc):
    """Извлекает текст из всех параграфов документа."""
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_tables(doc):
    """Извлекает таблицы из документа в формате JSON и текстовом виде."""
    tables_data = []
    text_tables = []
    for table in doc.tables:
        rows = []
        text_rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
            text_rows.append(" | ".join(row_data))  # Формируем строку текста из ячеек
        if rows:
            headers = rows[0]
            data_rows = rows[1:]
            tables_data.append([dict(zip(headers, row)) for row in data_rows])
            text_tables.append("\n".join(text_rows))  # Добавляем текстовую таблицу в список
    
    return tables_data, text_tables

# Открываем DOCX-файл и извлекаем данные
doc = docx.Document(docx_path)
data["text"] = extract_text(doc)
data["tables"], data["texttables"] = extract_tables(doc)

# Выводим JSON в stdout
sys.stdout.write(json.dumps(data, ensure_ascii=False, indent=4) + "\n")
